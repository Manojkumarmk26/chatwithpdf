import os
import logging
from typing import Union, Dict, List, Optional
from pathlib import Path
import fitz  # PyMuPDF
from PIL import Image
import docx
import numpy as np

logger = logging.getLogger(__name__)

class DocumentLoader:
    """
    A class to handle loading of various document types including PDF, DOCX, and images.
    """
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.tiff': 'image/tiff',
        '.bmp': 'image/bmp'
    }
    
    def __init__(self):
        """Initialize the DocumentLoader with default settings."""
        self.text_content = ""
        self.metadata = {}
    
    @classmethod
    def is_supported_file(cls, file_path: Union[str, Path]) -> bool:
        """Check if the file type is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            bool: True if file type is supported, False otherwise
        """
        ext = Path(file_path).suffix.lower()
        return ext in cls.SUPPORTED_EXTENSIONS
    
    def load(self, file_path: Union[str, Path]) -> Dict[str, str]:
        """Load and process the document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing 'text' and 'metadata' of the loaded document
            
        Raises:
            ValueError: If file type is not supported or file is corrupted
            FileNotFoundError: If file does not exist
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if not self.is_supported_file(file_path):
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        self.metadata = {
            'filename': file_path.name,
            'file_size': os.path.getsize(file_path),
            'file_type': self.SUPPORTED_EXTENSIONS.get(file_path.suffix.lower(), 'unknown')
        }
        
        try:
            if file_path.suffix.lower() == '.pdf':
                self._load_pdf(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                self._load_docx(file_path)
            elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                self._load_image(file_path)
            else:
                raise ValueError(f"Unhandled file type: {file_path.suffix}")
                
            return {
                'text': self.text_content,
                'metadata': self.metadata
            }
            
        except Exception as e:
            logger.error(f"Error loading file {file_path}: {str(e)}")
            raise ValueError(f"Failed to load file {file_path}: {str(e)}")
    
    def _load_pdf(self, file_path: Path):
        """Load and extract text from PDF file."""
        try:
            with fitz.open(file_path) as doc:
                self.metadata.update({
                    'page_count': len(doc),
                    'author': doc.metadata.get('author', ''),
                    'title': doc.metadata.get('title', ''),
                    'creation_date': doc.metadata.get('creationDate', '')
                })
                
                text_parts = []
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text("text")
                    text_parts.append(text)
                
                self.text_content = "\n\n".join(text_parts)
                
        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")
    
    def _load_docx(self, file_path: Path):
        """Load and extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            self.metadata.update({
                'paragraph_count': len(doc.paragraphs),
                'tables_count': len(doc.tables)
            })
            
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            self.text_content = "\n".join(text_parts)
            
        except Exception as e:
            raise ValueError(f"Error processing DOCX: {str(e)}")
    
    def _load_image(self, file_path: Path):
        """Load image file and prepare for OCR processing."""
        try:
            # Just store the image path for now, actual OCR will be handled by OCRProcessor
            with Image.open(file_path) as img:
                self.metadata.update({
                    'image_size': img.size,
                    'image_mode': img.mode
                })
            # Set empty content as OCR will be handled separately
            self.text_content = ""
            
        except Exception as e:
            raise ValueError(f"Error loading image: {str(e)}")
    
    def clear(self):
        """Clear the loaded content and metadata."""
        self.text_content = ""
        self.metadata = {}