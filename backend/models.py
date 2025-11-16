from pydantic import BaseModel, Field
from typing import List, Dict, Optional
from datetime import datetime

class FileMetadata(BaseModel):
    file_id: str
    filename: str
    file_type: str
    upload_date: datetime
    file_size: int
    chunk_count: int
    is_selected: bool = False

class DocumentChunk(BaseModel):
    chunk_id: str
    file_id: str
    filename: str
    content: str
    page_number: Optional[int] = None
    chunk_index: int
    is_table: bool = False

class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str
    timestamp: datetime
    source_files: List[str] = []

class ChatSession(BaseModel):
    session_id: str
    created_at: datetime
    messages: List[ChatMessage] = []
    selected_files: List[str] = []
    uploads: List[FileMetadata] = []

class QueryRequest(BaseModel):
    query: str
    session_id: str
    selected_file_ids: List[str]


class ExportRequest(BaseModel):
    session_id: str
    format: str  # "pdf", "docx", "txt"

# ============= backend/document_processor/loader.py =============
import fitz
from docx import Document
from PIL import Image
import io
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

class DocumentLoader:
    @staticmethod
    def load_pdf(file_path):
        """Load PDF and extract text + tables"""
        try:
            doc = fitz.open(file_path)
            content = {
                "text": "",
                "tables": [],
                "pages": len(doc)
            }
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                content["text"] += f"\n--- Page {page_num + 1} ---\n{text}"
                
                # Extract tables
                tables = page.find_tables()
                if tables:
                    for table in tables:
                        table_data = table.extract()
                        content["tables"].append({
                            "page": page_num + 1,
                            "data": table_data
                        })
            
            return content
        except Exception as e:
            logger.error(f"Error loading PDF: {e}")
            raise

    @staticmethod
    def load_docx(file_path):
        """Load DOCX and extract text + tables"""
        try:
            doc = Document(file_path)
            content = {"text": "", "tables": []}
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content["text"] += para.text + "\n"
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                content["tables"].append(table_data)
            
            return content
        except Exception as e:
            logger.error(f"Error loading DOCX: {e}")
            raise

    @staticmethod
    def load_image(file_path):
        """Load image for OCR processing"""
        try:
            image = Image.open(file_path)
            return {"image": image, "format": image.format}
        except Exception as e:
            logger.error(f"Error loading image: {e}")
            raise