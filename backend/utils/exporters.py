from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from docx import Document
import logging

logger = logging.getLogger(__name__)

class ChatExporter:
    @staticmethod
    def export_to_pdf(chat_session, output_path):
        """Export chat session to PDF"""
        try:
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            elements = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
            )
            elements.append(Paragraph("Chat Analysis Report", title_style))
            elements.append(Spacer(1, 12))
            
            # Files
            elements.append(Paragraph("Files Analyzed:", styles['Heading2']))
            for file in chat_session.uploads:
                elements.append(Paragraph(f"• {file.filename}", styles['Normal']))
            elements.append(Spacer(1, 20))
            
            # Messages
            elements.append(Paragraph("Chat History:", styles['Heading2']))
            for msg in chat_session.messages:
                role_style = 'Heading3' if msg.role == 'user' else 'Normal'
                elements.append(Paragraph(f"<b>{msg.role.upper()}:</b>", styles[role_style]))
                elements.append(Paragraph(msg.content, styles['Normal']))
                elements.append(Spacer(1, 12))
            
            doc.build(elements)
            logger.info(f"PDF exported to {output_path}")
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise

    @staticmethod
    def export_to_docx(chat_session, output_path):
        """Export chat session to Word"""
        try:
            doc = Document()
            doc.add_heading('Chat Analysis Report', 0)
            
            # Files
            doc.add_heading('Files Analyzed', level=1)
            for file in chat_session.uploads:
                doc.add_paragraph(file.filename, style='List Bullet')
            
            # Messages
            doc.add_heading('Chat History', level=1)
            for msg in chat_session.messages:
                doc.add_heading(msg.role.upper(), level=2)
                doc.add_paragraph(msg.content)
            
            doc.save(output_path)
            logger.info(f"DOCX exported to {output_path}")
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise

    @staticmethod
    def export_to_txt(chat_session, output_path):
        """Export chat session to TXT"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write("CHAT ANALYSIS REPORT\n")
                f.write("=" * 80 + "\n\n")
                
                f.write("FILES ANALYZED:\n")
                f.write("-" * 40 + "\n")
                for file in chat_session.uploads:
                    f.write(f"• {file.filename}\n")
                f.write("\n")
                
                f.write("CHAT HISTORY:\n")
                f.write("-" * 40 + "\n")
                for msg in chat_session.messages:
                    f.write(f"\n[{msg.role.upper()}]\n")
                    f.write(f"{msg.content}\n")
                    f.write("-" * 40 + "\n")
            
            logger.info(f"TXT exported to {output_path}")
        except Exception as e:
            logger.error(f"Error exporting to TXT: {e}")
            raise
